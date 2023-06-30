import os
import argparse
import glob
import html
import io
import re
import time
from pypdf import PdfReader, PdfWriter
from azure.identity import AzureDeveloperCliCredential
from azure.core.credentials import AzureKeyCredential
from azure.storage.blob import BlobServiceClient
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.indexes.models import *
from azure.search.documents import SearchClient
from azure.ai.formrecognizer import DocumentAnalysisClient

import requests
from msal import ConfidentialClientApplication
import openpyxl
from openpyxl.cell.text import RichText
from openpyxl.utils.datetime import from_excel
from docx import Document
import pandas as pd
from datetime import datetime
import json

from openpyxl.cell.cell import TYPE_FORMULA

MAX_SECTION_LENGTH = 1000
SENTENCE_SEARCH_LIMIT = 100
SECTION_OVERLAP = 100

parser = argparse.ArgumentParser(
    description="Prepare documents by extracting content from PDFs, splitting content into sections, uploading to blob storage, and indexing in a search index.",
    epilog="Example: prepdocs.py '..\data\*' --storageaccount myaccount --container mycontainer --searchservice mysearch --index myindex -v"
    )
parser.add_argument("files", help="Files to be processed")
parser.add_argument("--category", help="Value for the category field in the search index for all sections indexed in this run")
parser.add_argument("--skipblobs", action="store_true", help="Skip uploading individual pages to Azure Blob Storage")
parser.add_argument("--storageaccount", help="Azure Blob Storage account name")
parser.add_argument("--container", help="Azure Blob Storage container name")
parser.add_argument("--storagekey", required=False, help="Optional. Use this Azure Blob Storage account key instead of the current user identity to login (use az login to set current user for Azure)")
parser.add_argument("--tenantid", required=False, help="Optional. Use this to define the Azure directory where to authenticate)")
parser.add_argument("--searchservice", help="Name of the Azure Cognitive Search service where content should be indexed (must exist already)")
parser.add_argument("--index", help="Name of the Azure Cognitive Search index where content should be indexed (will be created if it doesn't exist)")
parser.add_argument("--searchkey", required=False, help="Optional. Use this Azure Cognitive Search account key instead of the current user identity to login (use az login to set current user for Azure)")
parser.add_argument("--remove", action="store_true", help="Remove references to this document from blob storage and the search index")
parser.add_argument("--removeall", action="store_true", help="Remove all blobs from blob storage and documents from the search index")
parser.add_argument("--localpdfparser", action="store_true", help="Use PyPdf local PDF parser (supports only digital PDFs) instead of Azure Form Recognizer service to extract text, tables and layout from the documents")
parser.add_argument("--formrecognizerservice", required=False, help="Optional. Name of the Azure Form Recognizer service which will be used to extract text, tables and layout from the documents (must exist already)")
parser.add_argument("--formrecognizerkey", required=False, help="Optional. Use this Azure Form Recognizer account key instead of the current user identity to login (use az login to set current user for Azure)")
parser.add_argument("--verbose", "-v", action="store_true", help="Verbose output")
args = parser.parse_args()
#print(args)
#args.removeall = True
#args.formrecognizerservice = 'NAME OF YOUR Azure Form Recognizer '
args.localpdfparser = True
#print(args)
# Use the current user identity to connect to Azure services unless a key is explicitly set for any of them
azd_credential = AzureDeveloperCliCredential() if args.tenantid == None else AzureDeveloperCliCredential(tenant_id=args.tenantid)
default_creds = azd_credential if args.searchkey == None or args.storagekey == None else None
search_creds = default_creds if args.searchkey == None else AzureKeyCredential(args.searchkey)
if not args.skipblobs:
    storage_creds = default_creds if args.storagekey == None else args.storagekey
if not args.localpdfparser:
    # check if Azure Form Recognizer credentials are provided
    print(args.formrecognizerservice)
    if args.formrecognizerservice == None:
        print("Error: Azure Form Recognizer service is not provided. Please provide formrecognizerservice or use --localpdfparser for local pypdf parser.")
        exit(1)
    formrecognizer_creds = default_creds if args.formrecognizerkey == None else AzureKeyCredential(args.formrecognizerkey)

def blob_name_from_file_page(filename, page = 0):
    #This function is used to generate the blob name from the filename and page number. 
    #For PDF files, it adds the page number to the file's name.
    if os.path.splitext(filename)[1].lower() == ".pdf":
        return os.path.splitext(os.path.basename(filename))[0] + f"-{page}" + ".pdf"
    else:
        return os.path.basename(filename)               
            
def upload_blobs(file_content, filename):
    #uploading the contents of the file to Azure Blob Storage. Each page of the document is converted into a blob and uploaded separately.
    blob_service = BlobServiceClient(account_url=f"https://{args.storageaccount}.blob.core.windows.net", credential=storage_creds)
    blob_container = blob_service.get_container_client(args.container)
    if not blob_container.exists():
        blob_container.create_container()

    # if file is PDF split into pages and upload each page as a separate blob
    if os.path.splitext(filename)[1].lower() == ".pdf":
        #reader = PdfReader(file_content)
        reader = PdfReader(io.BytesIO(file_content))

        pages = reader.pages
        for i in range(len(pages)):
            blob_name = blob_name_from_file_page(filename, i)
            if args.verbose: print(f"\tUploading blob for page {i} -> {blob_name}")
            f = io.BytesIO()
            writer = PdfWriter()
            writer.add_page(pages[i])
            writer.write(f)
            f.seek(0)
            blob_container.upload_blob(blob_name, f, overwrite=True)       
    else:
        blob_name = blob_name_from_file_page(filename)
        f = io.BytesIO(file_content)
        blob_container.upload_blob(blob_name, f, overwrite=True)

def remove_blobs(filename):
    if args.verbose: print(f"Removing blobs for '{filename or '<all>'}'")
    blob_service = BlobServiceClient(account_url=f"https://{args.storageaccount}.blob.core.windows.net", credential=storage_creds)
    blob_container = blob_service.get_container_client(args.container)
    if blob_container.exists():
        if filename == None:
            blobs = blob_container.list_blob_names()
        else:
            prefix = os.path.splitext(os.path.basename(filename))[0]
            blobs = filter(lambda b: re.match(f"{prefix}-\d+\.pdf", b), blob_container.list_blob_names(name_starts_with=os.path.splitext(os.path.basename(prefix))[0]))
        for b in blobs:
            if args.verbose: print(f"\tRemoving blob {b}")
            blob_container.delete_blob(b)

def table_to_html(table):
    table_html = "<table>"
    rows = [sorted([cell for cell in table.cells if cell.row_index == i], key=lambda cell: cell.column_index) for i in range(table.row_count)]
    for row_cells in rows:
        table_html += "<tr>"
        for cell in row_cells:
            tag = "th" if (cell.kind == "columnHeader" or cell.kind == "rowHeader") else "td"
            cell_spans = ""
            if cell.column_span > 1: cell_spans += f" colSpan={cell.column_span}"
            if cell.row_span > 1: cell_spans += f" rowSpan={cell.row_span}"
            table_html += f"<{tag}{cell_spans}>{html.escape(cell.content)}</{tag}>"
        table_html +="</tr>"
    table_html += "</table>"
    return table_html

def get_document_text(filename):
    offset = 0
    page_map = []
    file_extension = os.path.splitext(filename)[1].lower()    
    
    if file_extension == ".pdf":
       
        if args.localpdfparser:
            #reader = PdfReader(filename)
            reader = PdfReader(io.BytesIO(file_content))
            pages = reader.pages
            for page_num, p in enumerate(pages):
                page_text = p.extract_text()
                page_map.append((page_num, offset, page_text))
                offset += len(page_text)
        else:
           
            if args.verbose: print(f"Extracting text from '{filename}' using Azure Form Recognizer")
            print("https://{args.formrecognizerservice}.cognitiveservices.azure.com/")
            form_recognizer_client = DocumentAnalysisClient(endpoint=f"https://{args.formrecognizerservice}.cognitiveservices.azure.com/", credential=formrecognizer_creds, headers={"x-ms-useragent": "azure-search-chat-demo/1.0.0"})
            with open(filename, "rb") as f:
                poller = form_recognizer_client.begin_analyze_document("prebuilt-layout", document = f)
            form_recognizer_results = poller.result()

            for page_num, page in enumerate(form_recognizer_results.pages):
                tables_on_page = [table for table in form_recognizer_results.tables if table.bounding_regions[0].page_number == page_num + 1]

                # mark all positions of the table spans in the page
                page_offset = page.spans[0].offset
                page_length = page.spans[0].length
                table_chars = [-1]*page_length
                for table_id, table in enumerate(tables_on_page):
                    for span in table.spans:
                        # replace all table spans with "table_id" in table_chars array
                        for i in range(span.length):
                            idx = span.offset - page_offset + i
                            if idx >=0 and idx < page_length:
                                table_chars[idx] = table_id

                # build page text by replacing characters in table spans with table html
                page_text = ""
                added_tables = set()
                for idx, table_id in enumerate(table_chars):
                    if table_id == -1:
                        page_text += form_recognizer_results.content[page_offset + idx]
                    elif not table_id in added_tables:
                        page_text += table_to_html(tables_on_page[table_id])
                        added_tables.add(table_id)

                page_text += " "
                page_map.append((page_num, offset, page_text))
                offset += len(page_text)
                

    elif file_extension == ".xlsx":
        # Customize this based on your needs
        # in this POC, I only wanted to index worksheets with the name Sales,
        # incase where were more than one worksheet
        sheet_names = ["Sales"] 
        file_content_in_memory = io.BytesIO(file_content)
        workbook = openpyxl.load_workbook(file_content_in_memory)
        for sheet_name in sheet_names:
            if sheet_name in workbook.sheetnames:
                data = extract_text_from_excel_sheet(workbook[sheet_name])
                for i, row in enumerate(data):
                    json_row = json.dumps(row)
                    page_map.append((i, offset, json_row))
                    offset += len(json_row)


    elif file_extension == ".docx":
        file_content_in_memory = io.BytesIO(file_content)
        text = extract_text_from_docx(file_content_in_memory)
        page_map.extend((i, offset + m.start(), m.group()) for i, m in enumerate(re.finditer('.{1,1000}(?:(?<=\.\s)|$)', text, re.DOTALL)))
  
    return page_map

#This was updated to deal with Excel Files.  Savign a row into the index with rowname and column names. 
#this will give ChatGPT a better frame of reference.
def split_text(page_map):
    SENTENCE_ENDINGS = [".", "!", "?"]
    WORDS_BREAKS = [",", ";", ":", " ", "(", ")", "[", "]", "{", "}", "\t", "\n"]
    if args.verbose: print(f"Splitting '{filename}' into sections")

    def find_page(offset):
        l = len(page_map)
        for i in range(l - 1):
            if offset >= page_map[i][1] and offset < page_map[i + 1][1]:
                return i
        return l - 1
    
    if filename.lower().endswith('.xlsx'):
        # The content is a dictionary, not text
        # Assuming that each dictionary is a single row from the Excel file
       
        content_list = json.loads("[" + ", ".join(p[2] for p in page_map) + "]") # Parse the content into a list of dictionaries

        for i, row in enumerate(content_list):
            row_name = row.pop('RowName', None)  # Extract and remove the 'RowName' from the row
            if row_name is not None:  # Check if the RowName exists
                row_name = row_name.strip()  # Remove leading and trailing spaces
                if row_name == '':  # Check if the stripped RowName is an empty string
                    continue  # Skip rows with no name after stripping
            else:
                continue  # Skip rows with no name

            chunks = []  # List to hold the chunks of the row
            chunk = {'RowName': row_name}  # Dictionary to hold the current chunk

            for key, value in row.items():
                # If adding the next key-value pair would exceed the limit, yield the current chunk
                if len(json.dumps(chunk)) + len(json.dumps({key: value})) > MAX_SECTION_LENGTH:
                    chunks.append(chunk)
                    chunk = {'RowName': row_name}

                chunk[key] = value

            # Add the last chunk if it's not empty
            if chunk:
                chunks.append(chunk)

            for chunk in chunks:
                yield chunk, i

    else:

        all_text = "".join(p[2] for p in page_map)
        length = len(all_text)
        start = 0
        end = length
        while start + SECTION_OVERLAP < length:
            last_word = -1
            end = start + MAX_SECTION_LENGTH

            if end > length:
                end = length
            else:
                # Try to find the end of the sentence
                while end < length and (end - start - MAX_SECTION_LENGTH) < SENTENCE_SEARCH_LIMIT and all_text[end] not in SENTENCE_ENDINGS:
                    if all_text[end] in WORDS_BREAKS:
                        last_word = end
                    end += 1
                if end < length and all_text[end] not in SENTENCE_ENDINGS and last_word > 0:
                    end = last_word # Fall back to at least keeping a whole word
            if end < length:
                end += 1

            # Try to find the start of the sentence or at least a whole word boundary
            last_word = -1
            while start > 0 and start > end - MAX_SECTION_LENGTH - 2 * SENTENCE_SEARCH_LIMIT and all_text[start] not in SENTENCE_ENDINGS:
                if all_text[start] in WORDS_BREAKS:
                    last_word = start
                start -= 1
            if all_text[start] not in SENTENCE_ENDINGS and last_word > 0:
                start = last_word
            if start > 0:
                start += 1

            section_text = all_text[start:end]
            yield (section_text, find_page(start))

            last_table_start = section_text.rfind("<table")
            if (last_table_start > 2 * SENTENCE_SEARCH_LIMIT and last_table_start > section_text.rfind("</table")):
                # If the section ends with an unclosed table, we need to start the next section with the table.
                # If table starts inside SENTENCE_SEARCH_LIMIT, we ignore it, as that will cause an infinite loop for tables longer than MAX_SECTION_LENGTH
                # If last table starts inside SECTION_OVERLAP, keep overlapping
                if args.verbose: print(f"Section ends with unclosed table, starting next section with the table at page {find_page(start)} offset {start} table start {last_table_start}")
                start = min(end - SECTION_OVERLAP, start + last_table_start)
            else:
                start = end - SECTION_OVERLAP
            
        if start + SECTION_OVERLAP < end:
            yield (all_text[start:end], find_page(start))

def create_search_index():
    #This function is responsible for creating a new search index on Azure Cognitive Search service if it 
    # does not already exist. It first creates an instance of the SearchIndexClient with the provided 
    # endpoint and credentials. Then it checks if the index already exists. If it doesn't, it defines a 
    # new SearchIndex with the specified fields and settings, then creates it using the client's create_index() method.
    # The SearchIndex object includes a name and a list of fields. Each field is defined with its name, type, and other properties.
    # A SemanticSettings object is also included, which specifies configurations for semantic search capabilities. 
    # In this case, only the 'content' field is prioritized in the semantic configuration.
    # Adding contentType
    if args.verbose: print(f"Ensuring search index {args.index} exists")
    index_client = SearchIndexClient(endpoint=f"https://{args.searchservice}.search.windows.net/",
                                     credential=search_creds)
    if args.index not in index_client.list_index_names():
        index = SearchIndex(
            name=args.index,
            fields=[
                SimpleField(name="id", type="Edm.String", key=True),
                SearchableField(name="content", type="Edm.String", analyzer_name="en.microsoft"),
                SimpleField(name="category", type="Edm.String", filterable=True, facetable=True),
                SimpleField(name="sourcepage", type="Edm.String", filterable=True, facetable=True),
                SimpleField(name="sourcefile", type="Edm.String", filterable=True, facetable=True),
                SimpleField(name="contentType", type="Edm.String", filterable=True, facetable=True)
            ],
            semantic_settings=SemanticSettings(
                configurations=[SemanticConfiguration(
                    name='default',
                    prioritized_fields=PrioritizedFields(
                        title_field=None, prioritized_content_fields=[SemanticField(field_name='content')]))])
        )
        if args.verbose: print(f"Creating {args.index} search index")
        index_client.create_index(index)
    else:
        if args.verbose: print(f"Search index {args.index} already exists")

def index_sections(filename, sections):
    # This function is responsible for indexing the given sections into the search index. 
    # It first creates an instance of SearchClient with the provided endpoint, index name, and credentials. 
    # Then it iterates over the sections, adding them to a batch. When the batch size reaches 1000, 
    # it uploads the batch of documents to the index using the client's upload_documents() method, 
    # then clears the batch. It also prints the number of sections indexed and the number of successful indices 
    # if verbose output is enabled. If there are any remaining sections in the batch after iterating over all sections, it uploads these as well.
    if args.verbose: print(f"Indexing sections from '{filename}' into search index '{args.index}'")
    search_client = SearchClient(endpoint=f"https://{args.searchservice}.search.windows.net/",
                                    index_name=args.index,
                                    credential=search_creds)
    i = 0
    batch = []
    for s in sections:
        batch.append(s)
        i += 1
        if i % 1000 == 0:
            results = search_client.upload_documents(documents=batch)
            succeeded = sum([1 for r in results if r.succeeded])
            if args.verbose: print(f"\tIndexed {len(results)} sections, {succeeded} succeeded")
            batch = []

    if len(batch) > 0:
        results = search_client.upload_documents(documents=batch)
        succeeded = sum([1 for r in results if r.succeeded])
        if args.verbose: print(f"\tIndexed {len(results)} sections, {succeeded} succeeded")

def remove_from_index(filename):
    if args.verbose: print(f"Removing sections from '{filename or '<all>'}' from search index '{args.index}'")
    search_client = SearchClient(endpoint=f"https://{args.searchservice}.search.windows.net/",
                                    index_name=args.index,
                                    credential=search_creds)
    while True:
        filter = None if filename == None else f"sourcefile eq '{os.path.basename(filename)}'"
        r = search_client.search("", filter=filter, top=1000, include_total_count=True)
        if r.get_count() == 0:
            break
        r = search_client.delete_documents(documents=[{ "id": d["id"] } for d in r])
        if args.verbose: print(f"\tRemoved {len(r)} sections from index")
        # It can take a few seconds for search results to reflect changes, so wait a bit
        time.sleep(2)       


def download_file_content(accessToken, siteId, itemId):   
    fileDownloadUrl = f"https://graph.microsoft.com/v1.0/sites/{siteId}/lists/companies/items/{itemId}/driveItem/content";
    headers = { 'Authorization': f'Bearer {accessToken}' }
    response = requests.get(fileDownloadUrl, headers=headers, stream=True)
    #print(response.content)
    response.raise_for_status()  # Raises stored HTTPError, if one occurred
    return response.content

#concatentating rowName -  assuming the first cell in the row is the row name
def extract_text_from_excel_sheet(worksheet):
    rows = list(worksheet.iter_rows())
    headers = [cell.value for cell in rows[0]]  # assuming the first row contains headers

    data = []
    for row in rows[1:]:
        row_dict = {"RowName": row[0].value}  # assuming the first cell in the row is the row name
        for header, cell in zip(headers[1:], row[1:]):  # start from the second cell
            value = cell.value
            if cell.data_type == TYPE_FORMULA:
                value = cell.value
            if isinstance(value, str):
                row_dict[header] = value
            elif isinstance(value, RichText):
                row_dict[header] = ' '.join(run.text for run in value)
            elif isinstance(value, (int, float, bool, complex)):
                row_dict[header] = str(value)
            elif isinstance(value, datetime):
                row_dict[header] = value.strftime('%Y-%m-%d %H:%M:%S')
            elif cell.is_date:
                date_value = from_excel(value)
                row_dict[header] = str(date_value)
            elif cell.data_type == 'e':
                row_dict[header] = 'ERROR'
        data.append(row_dict)
    return data

def extract_text_from_docx(file_content):
    document = Document(file_content)
    text = '\n'.join([paragraph.text for paragraph in document.paragraphs])
    return text

#WE WILL MOST LIKELY ADD A DATE FIELD AND HAVE AZURE SEARCH RETRIEVE DATA BASED ON DATE
#PLEASE ADD content type in your index.  
#Adding ContentType so excel can be better handled
def create_sections(filename, page_map):
    for i, (section, pagenum) in enumerate(split_text(page_map)):
        
        if filename.lower().endswith('.xlsx'):
            content_type = 'dict'
        else:
            content_type = 'text'            
        yield {
            "id": re.sub("[^0-9a-zA-Z_-]","_",f"{filename}-{i}"),           
            "content": json.dumps(section) if isinstance(section, dict) else section,
            "category": args.category,
            "sourcepage": blob_name_from_file_page(filename, pagenum),
            "sourcefile": filename,
            "contentType": content_type
        }

#THESE SHOULD BE MOVED TO ENV OR AZURE VAULT.  THEY ARE HERE FOR CONVENIENCE ONLY.  
CLIENT_ID = 'YOUR AZURE APP CLIENT_ID';
TENANT_ID = 'YOUR AZURE TENANT_ID';
CLIENT_SECRET = 'YOUR AZURE APP CLIENT SECRET';
#THIS IS THE SHAREOINT SITE_ID...IN THIS FORMAT  
SITE_ID = 'COMPANY.sharepoint.com,BLAHBLAHBLAH';
#THIS IS THE NAME OF THE DOCUMENT LIBRARY.  I AM ONLY TESTING IN ONE DOCUMENT LIBRARY (NOT THE GENERAL).  
#YOU CAN LOOP THROUGH YOUR SITE.  FOR THIS POC, I ONLY FOCUSED ON PARTICULAR ONE
LIBRARY_ID = 'NAME OF THE DOCUMENT LIBRARY';

def get_access_token(tenant_id, client_id, client_secret):
    authority = f"https://login.microsoftonline.com/{tenant_id}"
    app = ConfidentialClientApplication(
        client_id,
        authority=authority,
        client_credential=client_secret,
    )
    result = None
    result = app.acquire_token_silent(["https://graph.microsoft.com/.default"], account=None)

    if not result:
        result = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])

    if "access_token" in result:
        return result["access_token"]
    else:
        print(result.get("error"))
        print(result.get("error_description"))
        print(result.get("correlation_id"))  # You may need this when reporting a bug
        return None
    
def get_files_in_library(accessToken, siteId, libraryId):
    url = f"https://graph.microsoft.com/v1.0/sites/{siteId}/lists/{libraryId}/items"
    headers = {'Authorization': f'Bearer {accessToken}', 'Accept': 'application/json'}

    response = requests.get(url, headers=headers)
    response.raise_for_status()  # Raises stored HTTPError, if one occurred

    items = response.json().get('value', [])
    file_items = []

    for item in items:
        fileDetailsUrl = f"https://graph.microsoft.com/v1.0/sites/{siteId}/lists/{libraryId}/items/{item['id']}?expand=fields"
        fileDetailsResponse = requests.get(fileDetailsUrl, headers=headers)
        fileDetailsResponse.raise_for_status()  # Raises stored HTTPError, if one occurred
        fileDetails = fileDetailsResponse.json()        
       
        #Customize this for yourself.  
        #I'm only uploading documents that are marked for IndexYN == YES
        #THIS ABILITY, YOU CANNOT DO WITH SHAREPOINT INDEXER
        if fileDetails['fields'].get('IndexYN') == "Yes" and  fileDetails['fields'].get('ContentType') == "Document":
            file_items.append(fileDetails)
        
        #print(file_items)
    return file_items


#code starts here

if args.removeall:
    remove_blobs(None)
    remove_from_index(None)
else:
    if not args.remove:
        create_search_index()
    
    print(f"Processing files...")    
    
    # Get access token
    token = get_access_token(TENANT_ID, CLIENT_ID, CLIENT_SECRET)
    
    # Get files in SharePoint Document Library
    fileDetails = get_files_in_library(token, SITE_ID, LIBRARY_ID)
    
    for fileDetail in fileDetails:
        filename = fileDetail['fields']['FileLeafRef']
        if args.verbose: print(f"Processing '{filename}'")
        
        # Download the file from SharePoint directly into memory
        file_content = download_file_content(token, SITE_ID, fileDetail['id'])
        
        if args.remove:
            remove_blobs(filename)
            remove_from_index(filename)
        elif args.removeall:
            remove_blobs(None)
            remove_from_index(None)
        else:
            if not args.skipblobs:
                # You need to download the file from SharePoint and upload it to Blob Storage   
                print(filename);             
                upload_blobs(file_content, filename)
            # You need to get the document text from the downloaded file
            page_map = get_document_text(filename)
            sections = create_sections(os.path.basename(filename), page_map)
            index_sections(os.path.basename(filename), sections)



